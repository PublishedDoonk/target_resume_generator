from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
import streamlit as st
from io import BytesIO
import re

def add_bulleted_paragraph(doc: Document, text: str, level: int = 0):
    paragraph = doc.add_paragraph(style='ListBullet')
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.space_before = Pt(0)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph_format.left_indent = Pt(18 * level)
    run = paragraph.add_run(text)

@st.cache(allow_output_mutation=True)
def export_to_word(title: str, user_data: dict) -> None:
    """Export the user's generated target_resume to a Word document."""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    
    paragraph = doc.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    name_run = paragraph.add_run(user_data['personal_info']['name'] + '\n')
    name_run.font.size = Pt(14)
    name_run.bold = True
    
    bio_info = f"""{user_data['personal_info']['street']}
{user_data['personal_info']['city']}, {user_data['personal_info']['state']} {user_data['personal_info']['zip_code']}
Phone: {user_data['personal_info']['phone']}
Email: {user_data['personal_info']['email']}
{user_data['personal_info']['linkedin']}
{user_data['personal_info']['github']}
{user_data['personal_info']['website']}
    """.strip()
    bio_run = paragraph.add_run(bio_info)
    
    paragraph2 = doc.add_paragraph()
    paragraph2_format = paragraph2.paragraph_format
    paragraph2_format.space_after = Pt(0)
    paragraph2_format.space_before = Pt(23.1)
    paragraph2_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p2_header_run = paragraph2.add_run('SUMMARY: ')
    p2_header_run.font.size = Pt(12)
    p2_header_run.bold = True
    p2_summary_run = paragraph2.add_run(user_data['personal_info']['summary'])
    
    education_header = doc.add_paragraph()
    education_header_format = education_header.paragraph_format
    education_header_format.space_after = Pt(0)
    education_header_format.space_before = Pt(24.2)
    education_header_run = education_header.add_run('EDUCATION')
    education_header_run.font.size = Pt(12)
    education_header_run.bold = True
    
    schools = user_data['educational_info']['schools']
    for school in schools:
        school_paragraph = doc.add_paragraph()
        school_paragraph_format = school_paragraph.paragraph_format
        school_paragraph_format.space_after = Pt(0)
        school_paragraph_format.space_before = Pt(9.75)
        school_paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        school_info = f"""{school['school']}
{school['degree']}, {school['grad_date']}
{school['major']}
{school['gpa']}"""
        school_run = school_paragraph.add_run(school_info)
    
    jobs_header = doc.add_paragraph()
    jobs_header_format = jobs_header.paragraph_format
    jobs_header_format.space_after = Pt(0)
    jobs_header_format.space_before = Pt(24.2)
    jobs_header_run = jobs_header.add_run('PROFESSIONAL EXPERIENCE')
    jobs_header_run.font.size = Pt(12)
    jobs_header_run.bold = True
    
    jobs = user_data['work_info']['jobs']
    bullets = user_data['target_resume']['bullets']
    
    for job, bullet in zip(jobs, bullets):
        job_paragraph = doc.add_paragraph()
        job_paragraph_format = job_paragraph.paragraph_format
        job_paragraph_format.space_after = Pt(6.9)
        job_paragraph_format.space_before = Pt(9.75)
        job_paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        job_info = f"""{job['start_date']} - {job['end_date']}
{job['position']}
{job['company name']}, {job['company city']}, {job['company state']}"""
        job_header_run = job_paragraph.add_run(job_info)
        for b in bullet:
            add_bulleted_paragraph(doc, b, 1)
            
    skills_header = doc.add_paragraph()
    skills_header_format = skills_header.paragraph_format
    skills_header_format.space_after = Pt(6.9)
    skills_header_format.space_before = Pt(24.2)
    skills_header_run = skills_header.add_run('SKILLS')
    skills_header_run.font.size = Pt(12)
    skills_header_run.bold = True
    
    skills = user_data['target_resume']['skills']
    #skills = skills.split('\n')
    
    for skill in skills:
        add_bulleted_paragraph(doc, skill.split('\t')[-1], 1)
        
    achievements = user_data['achievements']
    if achievements:
        achievements = [a.split('\t')[-1] for a in achievements.split('\n')]
        
        achievements_header = doc.add_paragraph()
        achievements_header_format = achievements_header.paragraph_format
        achievements_header_format.space_after = Pt(6.9)
        achievements_header_format.space_before = Pt(24.2)
        achievements_header_run = achievements_header.add_run('ACHIEVEMENTS')
        achievements_header_run.font.size = Pt(12)
        achievements_header_run.bold = True
        
        for achievement in achievements:
            add_bulleted_paragraph(doc, achievement, 1)
    
    document_obj: BytesIO = BytesIO()
    
    doc.save(document_obj)
    document_obj.seek(0)
    
    return document_obj
    
    
